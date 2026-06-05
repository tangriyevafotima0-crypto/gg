"""Export functionality for PDF, Word, Markdown, and TXT formats."""

import io
import tempfile
import os

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt


def export_as_pdf(doc: fitz.Document) -> bytes:
    """Export the document as PDF bytes."""
    return doc.tobytes()


def export_as_docx(doc: fitz.Document) -> bytes:
    """Export the document text as a Word (.docx) file."""
    docx_doc = DocxDocument()

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        if page_num > 0:
            docx_doc.add_page_break()

        # Add page header
        docx_doc.add_heading(f"Sahifa {page_num + 1}", level=2)

        # Add text paragraphs
        for paragraph_text in text.split("\n"):
            if paragraph_text.strip():
                p = docx_doc.add_paragraph(paragraph_text)
                for run in p.runs:
                    run.font.size = Pt(11)

    # Save to bytes
    buffer = io.BytesIO()
    docx_doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_as_markdown(doc: fitz.Document) -> str:
    """Export the document text as Markdown."""
    lines = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        lines.append(f"## Sahifa {page_num + 1}\n")
        lines.append(text)
        lines.append("\n---\n")

    return "\n".join(lines)


def export_as_txt(doc: fitz.Document) -> str:
    """Export the document text as plain text."""
    lines = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        lines.append(f"=== Sahifa {page_num + 1} ===\n")
        lines.append(text)
        lines.append("\n")

    return "\n".join(lines)
